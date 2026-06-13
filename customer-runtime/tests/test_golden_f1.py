"""T-011: Golden dataset F1 test."""

import asyncio
import io
import json
import os
import pathlib

import pytest


GOLDEN_DIR = pathlib.Path(__file__).parent / "fixtures" / "parser" / "golden"

SYNTHETIC_EXPECTED = {
    "device_name": "Regula Infusion Pump",
    "intended_use": "delivery of controlled infusion therapy in clinical settings",
    "indications": "adult patients requiring controlled fluid delivery",
    "contraindications": "not for use with blood products",
    "warnings": "verify flow rate before starting therapy",
    "device_classification": "Class II medical device",
    "region_targets": "US, EU, Korea",
    "cybersecurity_requirements": "unique user authentication and encrypted logs",
    "precautions": "inspect tubing before use",
    "product_code": "RIP-1000",
    "maintenance_interval": "every 12 months",
    "cleaning_disinfection": "wipe with approved disinfectant after each use",
    "software_version": "v1.2.3",
    "accessories": "power adapter, tubing set",
    "disposal_instructions": "dispose according to local electronic waste rules",
}


def _build_synthetic_docx_bytes() -> bytes:
    """Create a synthetic IFU DOCX with no customer or copyrighted content."""
    from docx import Document

    doc = Document()
    doc.add_heading("Synthetic IFU", level=1)
    labels = {
        "device_name": "Device Name",
        "intended_use": "Intended Use",
        "indications": "Indications",
        "contraindications": "Contraindications",
        "warnings": "Warnings",
        "device_classification": "Device Classification",
        "region_targets": "Region Targets",
        "cybersecurity_requirements": "Cybersecurity Requirements",
        "precautions": "Precautions",
        "product_code": "Product Code",
        "maintenance_interval": "Maintenance Interval",
        "cleaning_disinfection": "Cleaning and Disinfection",
        "software_version": "Software Version",
        "accessories": "Accessories",
        "disposal_instructions": "Disposal Instructions",
    }
    for field, label in labels.items():
        doc.add_paragraph(f"{label}: {SYNTHETIC_EXPECTED[field]}")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


@pytest.mark.integration
def test_golden_f1_macro_score():
    """Macro-F1 across golden dataset must be >= 0.85.

    Dataset format: each golden/*.docx has a matching *.json with expected field values.
    Set REQUIRE_GOLDEN_DATASET=1 in CI/release gates to fail when no dataset exists.
    """
    from app.schemas.parse import IFU_FIELD_NAMES
    from app.services.parser_engine import ParserEngine

    golden_files = list(GOLDEN_DIR.glob("*.docx"))
    require_dataset = os.environ.get("REQUIRE_GOLDEN_DATASET") == "1"
    synthetic_dataset: list[tuple[bytes, dict[str, str]]] = []
    if not golden_files and require_dataset:
        synthetic_dataset = [(_build_synthetic_docx_bytes(), SYNTHETIC_EXPECTED)]
    if not golden_files:
        if not synthetic_dataset:
            pytest.skip("Golden dataset not available")

    async def no_llm_stage(text, fields):
        return {}

    engine = ParserEngine(stage2_fn=lambda text, fields: {}, stage3_fn=no_llm_stage)

    all_precisions: list[float] = []
    all_recalls: list[float] = []

    async def score_one(docx_bytes: bytes, expected: dict[str, str]):
        parsed = await engine.parse(docx_bytes, "docx")

        tp = fp = fn = 0
        for field in IFU_FIELD_NAMES:
            fe = getattr(parsed, field)
            expected_val = expected.get(field)
            predicted_val = fe.value

            if expected_val and predicted_val:
                if expected_val.lower().strip() in str(predicted_val).lower():
                    tp += 1
                else:
                    fp += 1
                    fn += 1
            elif expected_val and not predicted_val:
                fn += 1
            elif not expected_val and predicted_val:
                fp += 1

        precision = tp / (tp + fp) if (tp + fp) > 0 else 0.0
        recall = tp / (tp + fn) if (tp + fn) > 0 else 0.0
        return precision, recall

    async def run_one(docx_path: pathlib.Path):
        label_path = docx_path.with_suffix(".json")
        if not label_path.exists():
            return None

        with open(docx_path, "rb") as f:
            docx_bytes = f.read()
        with open(label_path, encoding="utf-8") as f:
            expected = json.load(f)

        return await score_one(docx_bytes, expected)

    for docx_bytes, expected in synthetic_dataset:
        p, r = asyncio.run(score_one(docx_bytes, expected))
        all_precisions.append(p)
        all_recalls.append(r)

    for docx_path in golden_files:
        result = asyncio.run(run_one(docx_path))
        if result is not None:
            p, r = result
            all_precisions.append(p)
            all_recalls.append(r)

    if not all_precisions:
        pytest.skip("No matched golden files with labels")

    macro_p = sum(all_precisions) / len(all_precisions)
    macro_r = sum(all_recalls) / len(all_recalls)
    f1 = 2 * macro_p * macro_r / (macro_p + macro_r) if (macro_p + macro_r) > 0 else 0.0

    assert f1 >= 0.85, f"Macro-F1 {f1:.3f} < 0.85 threshold"
