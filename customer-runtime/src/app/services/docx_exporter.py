"""DOCX/JSON export service — SPEC-AUTHORING-001."""
import logging
from typing import Any

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.models.authoring_section_entry import AuthoringSectionEntry
from app.models.authoring_session import AuthoringSession

logger = logging.getLogger(__name__)


async def export_session(
    session_id: str,
    fmt: str,
    db: AsyncSession,
    template_sections: list[dict] | None = None,
) -> bytes | dict[str, Any]:
    """Export an authoring session as DOCX or JSON.

    Args:
        session_id: The session to export.
        fmt: "docx" or "json".
        db: Async database session.
        template_sections: Optional section metadata for ordering/titles.

    Returns:
        DOCX bytes for fmt="docx", dict for fmt="json".

    Raises:
        ValueError: For unsupported format.
        LookupError: If session not found.
    """
    if fmt not in ("docx", "json"):
        raise ValueError(f"Unsupported export format: {fmt!r}. Use 'docx' or 'json'.")

    result = await db.execute(
        select(AuthoringSession)
        .options(selectinload(AuthoringSession.entries))
        .where(AuthoringSession.session_id == session_id)
    )
    session = result.scalar_one_or_none()
    if session is None:
        raise LookupError(f"Session '{session_id}' not found.")

    # Build section index from template_sections if provided
    section_meta: dict[str, dict] = {}
    if template_sections:
        for sec in template_sections:
            section_meta[sec["section_id"]] = sec

    # Sort entries by sort_order if available, then by section_id
    def sort_key(entry: AuthoringSectionEntry) -> int:
        if entry.section_id in section_meta:
            return section_meta[entry.section_id].get("sort_order", 999)
        return 999

    sorted_entries = sorted(session.entries, key=sort_key)

    if fmt == "json":
        sections_data = []
        for entry in sorted_entries:
            meta = section_meta.get(entry.section_id, {})
            sections_data.append({
                "section_id": entry.section_id,
                "title": meta.get("title", entry.section_id),
                "sort_order": meta.get("sort_order", 0),
                "status": entry.status,
                "content": entry.content,
                "ai_draft": entry.ai_draft,
            })
        return {
            "session_id": session.session_id,
            "pack_id": session.pack_id,
            "status": session.status,
            "sections": sections_data,
        }

    # fmt == "docx"
    return _build_docx(session, sorted_entries, section_meta)


def _build_docx(
    session: AuthoringSession,
    entries: list[AuthoringSectionEntry],
    section_meta: dict[str, dict],
) -> bytes:
    """Build DOCX bytes from session entries."""
    from docx import Document  # type: ignore[import-untyped]
    from io import BytesIO

    doc = Document()
    doc.add_heading(f"Authoring Session Export — {session.pack_id}", level=1)
    doc.add_paragraph(f"Session ID: {session.session_id}")
    doc.add_paragraph(f"Status: {session.status}")
    doc.add_paragraph("")

    for entry in entries:
        meta = section_meta.get(entry.section_id, {})
        title = meta.get("title", entry.section_id)
        doc.add_heading(title, level=2)
        content = entry.content or entry.ai_draft or "(no content)"
        doc.add_paragraph(content)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
